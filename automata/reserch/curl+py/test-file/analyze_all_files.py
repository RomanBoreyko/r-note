import os
from transformers import pipeline
from docx import Document
import PyPDF2
import markdown
import logging

# Инициализация логирования
logging.basicConfig(level=logging.INFO)

# Инициализация модели для анализа настроений
classifier = pipeline("sentiment-analysis")

# Функция для чтения текстовых файлов
def read_text_file(file_path):
    with open(file_path, 'r', encoding='utf-8') as file:
        return file.read()

# Функция для чтения файлов .docx
def read_docx_file(file_path):
    doc = Document(file_path)
    full_text = []
    for para in doc.paragraphs:
        full_text.append(para.text)
    return '\n'.join(full_text)

# Функция для чтения PDF файлов
def read_pdf_file(file_path):
    reader = PyPDF2.PdfFileReader(file_path)
    full_text = []
    for page_num in range(reader.getNumPages()):
        page = reader.getPage(page_num)
        full_text.append(page.extract_text())
    return '\n'.join(full_text)

# Функция для чтения файлов Markdown
def read_md_file(file_path):
    with open(file_path, 'r', encoding='utf-8') as file:
        text = file.read()
    html = markdown.markdown(text)
    return html

# Основная функция для анализа всех файлов в указанной директории
def analyze_files_in_directory(directory):
    for root, dirs, files in os.walk(directory):
        logging.info(f"Поиск в директории: {root}")
        for file in files:
            file_path = os.path.join(root, file)
            logging.info(f"Найден файл: {file_path}")
            if file.endswith('.txt'):
                text = read_text_file(file_path)
            elif file.endswith('.docx'):
                text = read_docx_file(file_path)
            elif file.endswith('.pdf'):
                text = read_pdf_file(file_path)
            elif file.endswith('.md'):
                text = read_md_file(file_path)
            else:
                logging.warning(f"Неподдерживаемый формат файла: {file}")
                continue
            
            # Обработка текста, если он слишком длинный
            if len(text.split()) > 512:
                # Разбиваем текст на части по 512 токенов
                text_parts = [text[i:i+512] for i in range(0, len(text), 512)]
                results = []
                for part in text_parts:
                    results.append(classifier(part))
                result = combine_results(results)
            else:
                result = classifier(text)
            
            print(f"Файл: {file_path}")
            print(result)
            print()

# Функция для объединения результатов из разных частей текста
def combine_results(results):
    combined_result = {
        'label': 'POSITIVE' if all(result['label'] == 'POSITIVE' for result in results) else 'NEGATIVE',
        'score': sum(result['score'] for result in results) / len(results)
    }
    return [combined_result]

# Укажите путь к вашей директории
directory = "/home/orangepi/r_/r-note"
analyze_files_in_directory(directory)

# Функция для чтения текстовых файлов
def read_text_file(file_path):
    with open(file_path, 'r', encoding='utf-8') as file:
        return file.read()

# Функция для чтения .docx файлов
def read_docx_file(file_path):
    doc = Document(file_path)
    full_text = []
    for para in doc.paragraphs:
        full_text.append(para.text)
    return '\n'.join(full_text)

# Функция для чтения PDF файлов
def read_pdf_file(file_path):
    reader = PyPDF2.PdfFileReader(file_path)
    full_text = []
    for page_num in range(reader.getNumPages()):
        page = reader.getPage(page_num)
        full_text.append(page.extract_text())
    return '\n'.join(full_text)

# Функция для чтения Markdown файлов
def read_md_file(file_path):
    with open(file_path, 'r', encoding='utf-8') as file:
        text = file.read()
    html = markdown.markdown(text)
    return html


import os
from transformers import pipeline
from docx import Document
import PyPDF2
import markdown
import logging

# Инициализация логирования
logging.basicConfig(level=logging.INFO)

# Инициализация модели
classifier = pipeline("sentiment-analysis")

# Функции для чтения различных типов файлов
def read_text_file(file_path):
    with open(file_path, 'r', encoding='utf-8') as file:
        return file.read()

def read_docx_file(file_path):
    doc = Document(file_path)
    full_text = []
    for para in doc.paragraphs:
        full_text.append(para.text)
    return '\n'.join(full_text)

def read_pdf_file(file_path):
    reader = PyPDF2.PdfFileReader(file_path)
    full_text = []
    for page_num in range(reader.getNumPages()):
        page = reader.getPage(page_num)
        full_text.append(page.extract_text())
    return '\n'.join(full_text)

def read_md_file(file_path):
    with open(file_path, 'r', encoding='utf-8') as file:
        text = file.read()
    html = markdown.markdown(text)
    return html

# Функция для анализа всех файлов в директории
def analyze_files_in_directory(directory):
    results = []
    for root, dirs, files in os.walk(directory):
        logging.info(f"Поиск в директории: {root}")
        for file in files:
            file_path = os.path.join(root, file)
            logging.info(f"Найден файл: {file_path}")
            if file.endswith('.txt'):
                text = read_text_file(file_path)
            elif file.endswith('.docx'):
                text = read_docx_file(file_path)
            elif file.endswith('.pdf'):
                text = read_pdf_file(file_path)
            elif file.endswith('.md'):
                text = read_md_file(file_path)
            else:
                logging.warning(f"Неподдерживаемый формат файла: {file}")
                continue
            
            result = classifier(text)
            results.append(result)
            logging.info(f"Результаты анализа для файла {file_path}: {result}")
    
    return results

# Функция для объединения результатов анализа
def combine_results(results):
    if all(result[0]['label'] == 'POSITIVE' for result in results):
        return {'label': 'POSITIVE'}
    else:
        return {'label': 'NEGATIVE'}

# Основная логика скрипта
if __name__ == "__main__":
    directory = "/home/orangepi/r_/r-note"
    results = analyze_files_in_directory(directory)
    final_result = combine_results(results)
    print("Итоговый результат анализа всех файлов:", final_result)


